"""
Document Processor Module
Extracts text from PDF and DOCX files with OCR fallback
"""
import os
import re
from typing import Optional
import PyPDF2
from docx import Document
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
import io


class DocumentProcessor:
    """Process and extract text from various document formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def process_document(self, file_path: str) -> dict:
        """
        Extract text from document
        Returns: dict with text, metadata, and processing info
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._process_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _process_pdf(self, pdf_path: str) -> dict:
        """Extract text from PDF with OCR fallback"""
        text = ""
        pages = 0
        extraction_method = "direct"
        
        try:
            # Try direct text extraction first
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        # Add page marker for better section detection
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            # If text is too short, likely scanned PDF - use OCR
            if len(text.strip()) < 100:
                text = self._ocr_pdf(pdf_path)
                extraction_method = "ocr"
            
            # Split into sections
            sections = self.split_into_sections(text)
                
        except Exception as e:
            # Fallback to OCR if direct extraction fails
            try:
                text = self._ocr_pdf(pdf_path)
                extraction_method = "ocr_fallback"
                sections = self.split_into_sections(text)
            except Exception as ocr_error:
                return {
                    "success": False,
                    "error": f"Failed to extract text: {str(e)}, OCR also failed: {str(ocr_error)}"
                }
        
        return {
            "success": True,
            "text": text.strip(),
            "pages": pages,
            "extraction_method": extraction_method,
            "file_type": "pdf",
            "sections": sections
        }
    
    def _ocr_pdf(self, pdf_path: str) -> str:
        """OCR-based text extraction for scanned PDFs"""
        text = ""
        try:
            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=300)
            
            # OCR each page
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image, lang='eng')
                text += f"\n--- Page {i+1} ---\n{page_text}\n"
                
        except Exception as e:
            raise Exception(f"OCR failed: {str(e)}")
        
        return text
    
    def _process_docx(self, docx_path: str) -> dict:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract paragraphs with better formatting
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n\n"
            
            # Extract tables
            for table in doc.tables:
                text += "\n[TABLE]\n"
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
                text += "[/TABLE]\n\n"
            
            # Split into sections
            sections = self.split_into_sections(text)
            
            return {
                "success": True,
                "text": text.strip(),
                "pages": len(doc.paragraphs) // 20,  # Rough estimate
                "extraction_method": "docx",
                "file_type": "docx",
                "sections": sections
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to extract from DOCX: {str(e)}"
            }
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\:\;\-\(\)\[\]\"\'\/]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def split_into_sections(self, text: str) -> list:
        """Split document into logical sections with better detection"""
        sections = []
        
        # Try different section patterns (more comprehensive)
        patterns = [
            r'(?:SECTION|Section|Article|ARTICLE)\s+\d+[:\.\s]*([^\n]+)',
            r'(?:^|\n)(\d+\.\s+[A-Z][^\n]+)',
            r'(?:^|\n)([A-Z][A-Z\s&]{10,})(?:\n|:)',
            r'(?:^|\n)([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Clause|Agreement|Rights|Provision))'
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.MULTILINE)
            if matches:
                positions = [(m.start(), m.group(0)) for m in re.finditer(pattern, text, re.MULTILINE)]
                
                if len(positions) > 2:  # Found valid section markers
                    for i in range(len(positions)):
                        start = positions[i][0]
                        end = positions[i+1][0] if i < len(positions)-1 else len(text)
                        section_text = text[start:end].strip()
                        
                        if len(section_text) > 50:  # Minimum section length
                            sections.append({
                                'title': positions[i][1].strip(),
                                'text': section_text,
                                'position': start
                            })
                    return sections
        
        # Fallback: split by paragraphs
        paragraphs = text.split('\n\n')
        for i, para in enumerate(paragraphs):
            if len(para.strip()) > 50:
                sections.append({
                    'title': f'Paragraph {i+1}',
                    'text': para.strip(),
                    'position': text.find(para)
                })
        
        return sections
